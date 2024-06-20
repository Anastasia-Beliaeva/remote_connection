from docx import Document
from docx.shared import Pt
import pandas as pd
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx2pdf import convert

df = pd.read_excel('/Users/anastasiabelaeva/Desktop/Postgraduate/данные/CT/экономыш 2023/обратная связь/июнь 2024/db.xlsx')

# заменяю числовые значения на слова
df['CT_lvl'] = df['CT_lvl'].astype(str)
df['CT_lvl'] = df['CT_lvl'].str.replace('1', 'базовый')
df['CT_lvl'] = df['CT_lvl'].str.replace('2', 'высокий')
df['CT_lvl'] = df['CT_lvl'].str.replace('3', 'продвинутый')

df['critical_thinking_progress_positive'] = df['critical_thinking_progress_positive'].round()

# функция, которая пишет ворд, на вход принимает имя студента, его балл, уровень КМ,
# процент решенных заданий по проверки информации и процент решенных заданий по анализу и рефлексии

def document(name, points, CT_lvl, info_check, info_analys, time, progress):
    document = docx.Document()
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    heading = document.add_heading('Результаты теста, направленного на оценку критического мышления ', 1)
    heading.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_style = heading.style
    title_style.font.name = "Times New Roman"

    p = document.add_paragraph('Уважаемый респондент!')
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.style = document.styles['Normal']

    p = document.add_paragraph('Спасибо за участие в оценивании критического мышления. Тест, который Вы проходили, был частью исследования уровня критического мышления у студентов разных специальностей ведущих российских вузов. Сегодня считается, что критическое мышление может развиваться в двух ситуациях: как в результате специального обучения, так и благодаря накоплению практического жизненного опыта.',style='BodyText')
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p = document.add_paragraph('В тесте критическое мышление определялось как последовательность когнитивных действий, направленных на оценку качества исходной информации с целью определения проблемы, поиск возможных решений и выбор наилучшего из них, обоснование собственного вывода и выявление его ограничений и оценивались следующие умения:',style='BodyText')
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # document.add_heading('Heading, level 1', level=1)
    # document.add_paragraph('Intense quote', style='Intense Quote')

    document.add_paragraph(
        'Проверять информацию, включая группировку и ранжирование источников исходной информации, определять её актуальность и релевантность, оценивать компетентность и авторитетность источников информации;', style='List Number'
    )
    document.add_paragraph(
        'Анализировать и осмыслять информацию, на основе анализа информации выносить четкое суждение, разрабатывать истинные и валидные выводы и проводить рефлексию в отношении альтернативных объяснений.', style='List Number'
    )

    p = document.add_paragraph(
        'Эти две грани критического мышления могут подменять или дополнять друг друга. Поэтому, помимо общего уровня критического мышления обратите внимание на процент решенных заданий из каждой грани.',
        style='BodyText')
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p = document.add_paragraph('')
    p.add_run('Ваши результаты по тесту').bold = True

    #если списали (бытсро решили) - предупреждение
    if time == 1:
        p = document.add_paragraph(
            'Время выполнения теста составило меньше 20 минут, поэтому ваши результаты могут быть не точными.',
            style='BodyText')
        p.alignment = 0

    #балл за тест
    points = points
    p = document.add_paragraph('Балл за тест: ', style='BodyText')
    p.add_run(points)
    p.alignment = 0

    #прогресс с прошлого среза
    if progress > 0:
        p = document.add_paragraph(
            f'За год вы улучили свой результат на: {progress}.',
            style='BodyText')
        p.alignment = 0
    elif progress == 0:
        p = document.add_paragraph(f'За год ваш балл не изменился.', style='BodyText')
        p.alignment = 0
    else:
        pass
    # # уровень км
    CT_lvl = CT_lvl
    p = document.add_paragraph('Уровень критического мышления: ', style='BodyText')
    p.add_run(CT_lvl)
    p.alignment = 0


    # процент решенных заданий по проверки информации
    info_check = info_check
    p = document.add_paragraph('Процент решенных заданий на проверку информации: ', style='BodyText')
    p.add_run(info_check)
    p.alignment = 0
    # процент решенных заданий по анализу и рефлексии
    info_analys = info_analys
    p = document.add_paragraph('Процент решенных заданий на анализ и осмысление информации: ', style='BodyText')
    p.add_run(info_analys)
    p.alignment = 0

    document.add_page_break()
    # document.add_picture('monty-truth.png', width=Inches(1.25))
    #
    #функция, которая выделяет строки жирным шрифтом
    def make_rows_bold(*rows):
        for row in rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True


    records = (
        ('',
         'Вы можете: \n\nопределить основную мысль и ключевые термины статьи; \n\nоценить актуальность исходной информации.',
         '➕\n\nразличить факты и мнения, позитивные и нормативные суждения; \n\nуточнять термины; \n\nопределить компетентность и авторитетность источников.',
         '➕➕\n\nвыделять релевантную информацию; \n\nоценивать степень её непредвзятости.'),

        ('',
         'Вы можете: \n\nопределить проблему на основе текста (не тождественную основной мысли); \n\nвыбрать концепции и определить явные предположения анализа.',
         '➕\n\nвыявить причинно-следственные связи для построения прогнозов; \n\nсформулировать надёжные выводы на основе анализа; \n\nвыявить неявные предположения анализа.',
         '➕➕\n\nоценить ограничения анализа и реалистичность собственного вывода; \n\nвыявить противоречия анализа;\n\nоценить степень неопределённости выводов, полученных на основе анализа.')
    )
    # делаю табличку в будущем доке
    table = document.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = ''
    hdr_cells[1].text = 'Базовый'
    hdr_cells[2].text = 'Высокий'
    hdr_cells[3].text = 'Продвинутый'

    # делаю столбцы в табличке
    for Уровни, Базовый, Высокий, Продвинутый in records:
        row_cells = table.add_row().cells
        row_cells[0].text = Уровни
        row_cells[1].text = Базовый
        row_cells[2].text = Высокий
        row_cells[3].text = Продвинутый

    row = table.rows[0]
    Nombre_text_formatted = row.cells[0].paragraphs[0].add_run("Уровни")
    Nombre_text_formatted.bold = True
    row = table.rows[1]
    Nombre_text_formatted = row.cells[0].paragraphs[0].add_run("Проверка информации")
    Nombre_text_formatted.bold = True
    row = table.rows[2]
    Nombre_text_formatted = row.cells[0].paragraphs[0].add_run("Анализ и осмысление информации ")
    Nombre_text_formatted.bold = True

    make_rows_bold(table.rows[0])
    table.style = 'Table Grid'
    p = document.add_paragraph('')
    p.add_run('Спасибо за участие в исследовании!').bold = True

    # прописываю путь сохранения файла, в качестве именя файла будет почта студента
    path = ('/Users/anastasiabelaeva/Desktop/Postgraduate/данные/CT/экономыш 2023/обратная связь/июнь 2024/mails/'+uni+'/'+name+'.docx')
    document.save(path)

# функция, которая проходится по базе данных (эксельке) и создает документы для каждого студента
for index, row in df.iterrows():
    uni = row['university']
    name = row['id']
    points = str(row['critical_thinking24'])
    CT_lvl = row['CT_lvl']
    info_check = str(row['inf'])
    info_analys = str(row['an'])
    time = row['less_20min24']
    progress = row['critical_thinking_progress_positive']
    document(name, points, CT_lvl, info_check, info_analys, time, progress)



